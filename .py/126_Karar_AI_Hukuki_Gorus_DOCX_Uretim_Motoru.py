# -*- coding: utf-8 -*-

import sqlite3
import re
from pathlib import Path
from datetime import datetime
from difflib import SequenceMatcher
from collections import Counter

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DB_PATH = r"C:\Users\MSI\Desktop\kik_proje\.py\kik.db"
BASE_DIR = Path(r"C:\Users\MSI\Desktop\kik_proje")
RAPOR_DIR = BASE_DIR / "raporlar"
RAPOR_DIR.mkdir(exist_ok=True)

MIN_SKOR = 50
LIMIT = 3


def temizle(metin):
    if metin is None:
        return ""
    metin = str(metin).lower()
    metin = metin.replace("ı", "i").replace("ğ", "g").replace("ü", "u")
    metin = metin.replace("ş", "s").replace("ö", "o").replace("ç", "c")
    metin = re.sub(r"\s+", " ", metin)
    return metin.strip()


def dosya_adi_temizle(metin):
    metin = temizle(metin)
    metin = re.sub(r"[^a-zA-Z0-9]+", "_", metin)
    return metin.strip("_")[:80] or "hukuki_gorus"


def kelimeler(metin):
    return [k for k in re.findall(r"[a-zA-Z0-9]+", temizle(metin)) if len(k) > 2]


def konulari_ayir(sorgu):
    parcalar = re.split(r"\+|,| ve | ile | / ", sorgu)
    return [p.strip() for p in parcalar if p.strip()]


def benzerlik(a, b):
    return SequenceMatcher(None, temizle(a), temizle(b)).ratio()


def sonuc_tipi_bul(sonuc):
    s = temizle(sonuc)

    ret_kaliplari = [
        "kabul edilmemistir", "kabul edilmemis",
        "uygun bulunmamistir", "yerinde bulunmamistir",
        "yerinde gorulmemistir", "reddedilmistir",
        "reddedilmis", "ret edilmistir", "ret edilmis", "reddine"
    ]

    kabul_kaliplari = [
        "kabul edilmistir", "kabul edilmis",
        "uygun bulunmustur", "yerinde bulunmustur",
        "yerinde gorulmustur", "duzeltici islem"
    ]

    if "iptal" in s:
        return "İPTAL"
    if any(k in s for k in ret_kaliplari):
        return "RET"
    if any(k in s for k in kabul_kaliplari):
        return "KABUL"
    if "ret" in s or "redd" in s:
        return "RET"

    return "BELİRSİZ"


def kartlari_getir():
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    cur = con.cursor()

    rows = cur.execute("""
        SELECT
            id,
            karar_no,
            iddia_no,
            baslik,
            hukuki_soru,
            sonuc,
            emsal_ilke,
            anahtar_kelime,
            kurul_degerlendirmesi,
            mevzuat,
            guven
        FROM hukuki_kartlar
        ORDER BY karar_no, iddia_no
    """).fetchall()

    con.close()
    return rows


def konu_eslesir(konu, kart):
    konu_norm = temizle(konu)
    konu_kelimeleri = kelimeler(konu)

    alan = " ".join([
        temizle(kart["baslik"]),
        temizle(kart["hukuki_soru"]),
        temizle(kart["anahtar_kelime"]),
    ])

    if konu_norm in alan:
        return True

    if len(konu_kelimeleri) == 1:
        return konu_kelimeleri[0] in alan

    eslesen = sum(1 for k in konu_kelimeleri if k in alan)
    return eslesen >= max(2, len(konu_kelimeleri))


def skor_hesapla(konu, kart):
    skor = 0

    konu_norm = temizle(konu)
    konu_kelimeleri = kelimeler(konu)

    baslik = temizle(kart["baslik"])
    soru = temizle(kart["hukuki_soru"])
    sonuc = temizle(kart["sonuc"])
    ilke = temizle(kart["emsal_ilke"])
    anahtar = temizle(kart["anahtar_kelime"])

    if konu_norm == anahtar:
        skor += 120
    if konu_norm in anahtar:
        skor += 90
    if konu_norm in baslik:
        skor += 80
    if konu_norm in soru:
        skor += 60
    if konu_norm in ilke:
        skor += 30

    for k in konu_kelimeleri:
        if k in anahtar:
            skor += 15
        if k in baslik:
            skor += 12
        if k in soru:
            skor += 10
        if k in ilke:
            skor += 5
        if k in sonuc:
            skor += 2

    return skor


def tekrar_filtrele(sonuclar):
    secilen = []

    for skor, kart in sonuclar:
        tekrar = False

        for _, onceki in secilen:
            if benzerlik(kart["baslik"], onceki["baslik"]) >= 0.92:
                tekrar = True
                break
            if benzerlik(kart["emsal_ilke"], onceki["emsal_ilke"]) >= 0.92:
                tekrar = True
                break

        if not tekrar:
            secilen.append((skor, kart))

        if len(secilen) >= LIMIT:
            break

    return secilen


def konu_ara(konu):
    kartlar = kartlari_getir()
    sonuclar = []

    for kart in kartlar:
        if not konu_eslesir(konu, kart):
            continue

        skor = skor_hesapla(konu, kart)

        if skor >= MIN_SKOR:
            sonuclar.append((skor, kart))

    sonuclar.sort(key=lambda x: x[0], reverse=True)
    return tekrar_filtrele(sonuclar)


def p(doc, text="", bold=False, align=None, size=11):
    para = doc.add_paragraph()
    if align:
        para.alignment = align
    run = para.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return para


def heading(doc, text, level=1):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14 if level == 1 else 12)
    return para


def nihai_gorus_uret(dagilim):
    if not dagilim:
        return "Mevcut veri tabanında yeterli emsal bulunmadığından nihai ön görüş üretilememiştir."

    ana_tip = dagilim.most_common(1)[0][0]

    if ana_tip == "RET":
        return "Mevcut emsal kartlara göre başvurunun reddedilmesi yönünde ağırlıklı bir eğilim bulunmaktadır."
    if ana_tip == "KABUL":
        return "Mevcut emsal kartlara göre başvuru iddiasının kabulü yönünde ağırlıklı bir eğilim bulunmaktadır."
    if ana_tip == "İPTAL":
        return "Mevcut emsal kartlarda ihale sürecini etkileyen ve iptal sonucu doğurabilecek değerlendirmeler ağırlıktadır."

    return "Mevcut emsal kartlarda sonuç eğilimi net değildir; somut olayın ayrıca incelenmesi gerekir."


def docx_uret(sorgu):
    konular = konulari_ayir(sorgu)
    genel_dagilim = Counter()
    konu_sonuclari = {}

    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.3)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("KİK KARAR AI HUKUKİ GÖRÜŞ RAPORU")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    p(doc, f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "")

    heading(doc, "1. SORU")
    p(doc, sorgu)

    heading(doc, "2. AYRIŞTIRILAN KONULAR")
    for konu in konular:
        p(doc, f"- {konu}")

    heading(doc, "3. EMSAL KARARLAR")

    for konu in konular:
        sonuclar = konu_ara(konu)
        konu_sonuclari[konu] = sonuclar

        heading(doc, f"KONU: {konu}", 2)

        if not sonuclar:
            p(doc, "Bu konu için emsal bulunamadı.")
            continue

        for i, (skor, kart) in enumerate(sonuclar, start=1):
            tip = sonuc_tipi_bul(kart["sonuc"])
            genel_dagilim[tip] += 1

            p(doc, f"{i}) Karar No: {kart['karar_no']} | İddia No: {kart['iddia_no']} | Skor: {skor}", bold=True)
            p(doc, f"Sonuç Tipi: {tip}")
            p(doc, f"Hukuki Soru: {kart['hukuki_soru']}")
            p(doc, f"Sonuç: {kart['sonuc']}")
            p(doc, f"Emsal İlke: {kart['emsal_ilke']}")
            p(doc, "")

    heading(doc, "4. ORTAK HUKUKİ DEĞERLENDİRME")

    p(doc, "Mevcut emsal kartlar birlikte değerlendirildiğinde uyuşmazlığın tek bir başlık altında değil, birden fazla hukuki unsurun birlikte ele alınmasını gerektirdiği görülmektedir.")

    for konu in konular:
        sonuclar = konu_sonuclari.get(konu, [])
        if not sonuclar:
            p(doc, f"{konu} bakımından mevcut test veritabanında yeterli emsal bulunamamıştır.")
            continue

        skor, kart = sonuclar[0]
        p(doc, f"{konu} bakımından en güçlü emsal {kart['karar_no']} sayılı kararın {kart['iddia_no']} numaralı iddiasıdır. Bu emsale göre; {kart['emsal_ilke']}")

    heading(doc, "5. SONUÇ EĞİLİMİ")
    if not genel_dagilim:
        p(doc, "Sonuç eğilimi üretilemedi.")
    else:
        for tip, adet in genel_dagilim.most_common():
            p(doc, f"- {tip}: {adet} emsal")

    heading(doc, "6. NİHAİ ÖN GÖRÜŞ")
    p(doc, nihai_gorus_uret(genel_dagilim))

    heading(doc, "7. UYARI")
    p(doc, "Bu metin otomatik çoklu konu emsal analizine dayalı ön hukuki görüş taslağıdır. Nihai değerlendirme için karar tam metinleri, ihale dokümanı, başvuru dilekçesi ve ilgili mevzuat birlikte incelenmelidir.")

    tarih = datetime.now().strftime("%Y%m%d_%H%M%S")
    ad = dosya_adi_temizle(sorgu)
    dosya = RAPOR_DIR / f"126_hukuki_gorus_{ad}_{tarih}.docx"
    doc.save(dosya)

    return dosya


def main():
    print("=" * 80)
    print("126 - KİK KARAR AI HUKUKİ GÖRÜŞ DOCX ÜRETİM MOTORU")
    print("=" * 80)

    while True:
        sorgu = input("\nÇoklu konu / hukuki soru (çıkış=q): ").strip()

        if sorgu.lower() in ["q", "quit", "exit", "çıkış"]:
            break

        dosya = docx_uret(sorgu)

        print("\nDOCX RAPOR OLUŞTURULDU:")
        print(dosya)

    print("\nProgram sonlandırıldı.")


if __name__ == "__main__":
    main()